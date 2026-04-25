"""DOCX renderer — produces a styled resume from a ResumeIR model.

Styling is driven by ``TemplateConfig`` (with defaults matching the original
hardcoded values from generate_resume_docx.js).
"""

from __future__ import annotations

import re
from types import SimpleNamespace

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips
from docx.table import Table

from resume_builder.models.resume import ResumeIR
from resume_builder.models.template_config import TemplateConfig
from resume_builder.templates import (
    resolve_docx_font,
    resolve_docx_page,
    resolve_docx_style,
)

# JS sizes are in half-points; python-docx Pt() takes full points.
HP = 0.5


def _hex_to_rgb(hex_color: str) -> RGBColor:
    """Convert ``#RRGGBB`` to ``RGBColor``."""
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _hex_strip(hex_color: str) -> str:
    """Strip ``#`` prefix for OxmlElement attributes."""
    return hex_color.lstrip("#")


# ---------------------------------------------------------------------------
# Low-level OxmlElement helpers
# ---------------------------------------------------------------------------

def _set_bottom_border(paragraph, color: str, size: int) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _remove_cell_borders(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_cell_margins(cell, top: int, bottom: int, left: int, right: int) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for edge, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _set_cell_width(cell, width: int) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(width))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _add_tab_stop_right(paragraph, text_width: int) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(text_width))
    tab.set(qn("w:leader"), "none")
    tabs.append(tab)
    pPr.append(tabs)


def _configure_bullet_numbering(doc: Document) -> str:
    numbering_part = doc.part.numbering_part
    numbering_elm = numbering_part.numbering_definitions._numbering

    existing_abstract = numbering_elm.findall(qn("w:abstractNum"))
    next_abstract_id = str(len(existing_abstract) + 100)

    existing_nums = numbering_elm.findall(qn("w:num"))
    next_num_id = str(len(existing_nums) + 100)

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), next_abstract_id)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "hybridMultilevel")
    abstract_num.append(multi)

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    numFmt = OxmlElement("w:numFmt")
    numFmt.set(qn("w:val"), "bullet")
    lvl.append(numFmt)
    lvlText = OxmlElement("w:lvlText")
    lvlText.set(qn("w:val"), "•")
    lvl.append(lvlText)
    lvlJc = OxmlElement("w:lvlJc")
    lvlJc.set(qn("w:val"), "left")
    lvl.append(lvlJc)
    pPr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    pPr.append(ind)
    lvl.append(pPr)
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Symbol")
    rFonts.set(qn("w:hAnsi"), "Symbol")
    rFonts.set(qn("w:hint"), "default")
    rPr.append(rFonts)
    lvl.append(rPr)

    abstract_num.append(lvl)
    numbering_elm.insert(0, abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), next_num_id)
    abstractNumId_el = OxmlElement("w:abstractNumId")
    abstractNumId_el.set(qn("w:val"), next_abstract_id)
    num.append(abstractNumId_el)
    numbering_elm.append(num)

    return next_num_id


def _make_bullet_paragraph(paragraph, num_id: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numPr.append(ilvl)
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), num_id)
    numPr.append(numId_el)
    pPr.append(numPr)


# ---------------------------------------------------------------------------
# Text-run helpers
# ---------------------------------------------------------------------------

def _plain_run(paragraph, text: str, font_name: str, size_hp: int, color: RGBColor | None = None):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size_hp * HP)
    if color:
        run.font.color.rgb = color
    return run


def _bold_run(paragraph, text: str, font_name: str, size_hp: int, color: RGBColor | None = None):
    run = _plain_run(paragraph, text, font_name, size_hp, color)
    run.bold = True
    return run


def _italic_run(paragraph, text: str, font_name: str, size_hp: int, color: RGBColor | None = None):
    run = _plain_run(paragraph, text, font_name, size_hp, color)
    run.italic = True
    return run


def _bold_italic_run(paragraph, text: str, font_name: str, size_hp: int, color: RGBColor | None = None):
    run = _plain_run(paragraph, text, font_name, size_hp, color)
    run.bold = True
    run.italic = True
    return run


# ---------------------------------------------------------------------------
# Markdown-bold parser  (**bold** markers inside text)
# ---------------------------------------------------------------------------
_BOLD_RE = re.compile(r"(\*\*[^*]+\*\*)")


def _add_parsed_runs(paragraph, text: str, font_name: str, size_hp: int, color: RGBColor | None = None):
    parts = _BOLD_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            _bold_run(paragraph, part[2:-2], font_name, size_hp, color)
        else:
            _plain_run(paragraph, part, font_name, size_hp, color)


# ---------------------------------------------------------------------------
# Paragraph-level builders
# ---------------------------------------------------------------------------

def _set_spacing(paragraph, before: int | None = None, after: int | None = None):
    fmt = paragraph.paragraph_format
    if before is not None:
        fmt.space_before = Twips(before)
    if after is not None:
        fmt.space_after = Twips(after)


def _add_section_heading(doc: Document, text: str, ctx: SimpleNamespace) -> None:
    s = resolve_docx_style("section", ctx.config)
    p = doc.add_paragraph()
    _set_spacing(p, before=280, after=80)
    _set_bottom_border(p, _hex_strip(ctx.config.colors.rule), 6)
    _bold_run(p, text.upper(), s.font_name, s.size_hp, _hex_to_rgb(s.color_hex))


def _add_subheading(doc: Document, left: str, right: str | None, ctx: SimpleNamespace) -> None:
    s_company = resolve_docx_style("company", ctx.config)
    s_body = resolve_docx_style("body", ctx.config)
    p = doc.add_paragraph()
    _set_spacing(p, before=200, after=40)
    _bold_run(p, left, s_company.font_name, s_company.size_hp, _hex_to_rgb(s_company.color_hex))
    if right:
        _add_tab_stop_right(p, ctx.page.text_width)
        _plain_run(p, f"\t{right}", s_body.font_name, s_body.size_hp, _hex_to_rgb(ctx.config.colors.muted))


def _add_role_title(doc: Document, title: str, dates: str | None, ctx: SimpleNamespace) -> None:
    s = resolve_docx_style("role", ctx.config)
    p = doc.add_paragraph()
    _set_spacing(p, before=100, after=20)
    _bold_italic_run(p, title, s.font_name, s.size_hp, _hex_to_rgb(s.color_hex))
    if dates:
        _add_tab_stop_right(p, ctx.page.text_width)
        r = _plain_run(p, f"\t{dates}", s.font_name, s.size_hp, _hex_to_rgb(ctx.config.colors.muted))
        r.italic = True


def _add_role_description(doc: Document, text: str, ctx: SimpleNamespace) -> None:
    s = resolve_docx_style("role_desc", ctx.config)
    p = doc.add_paragraph()
    _set_spacing(p, before=0, after=40)
    _italic_run(p, text, s.font_name, s.size_hp, _hex_to_rgb(s.color_hex))


def _add_bullet(doc: Document, text: str, num_id: str, ctx: SimpleNamespace) -> None:
    s = resolve_docx_style("bullet", ctx.config)
    p = doc.add_paragraph()
    _set_spacing(p, before=40, after=40)
    _make_bullet_paragraph(p, num_id)
    _add_parsed_runs(p, text, s.font_name, s.size_hp)


def _format_bullet_text(bullet) -> str:
    if bullet.label:
        return f"**{bullet.label}:** {bullet.text}"
    return bullet.text


# ---------------------------------------------------------------------------
# Skills table
# ---------------------------------------------------------------------------

def _add_skills_table(doc: Document, skills, ctx: SimpleNamespace) -> None:
    s_label = resolve_docx_style("skill_label", ctx.config)
    s_value = resolve_docx_style("skill_value", ctx.config)
    light_bg = _hex_strip(ctx.config.colors.light_bg)

    table = doc.add_table(rows=len(skills), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_idx, skill in enumerate(skills):
        row = table.rows[row_idx]
        label_cell = row.cells[0]
        value_cell = row.cells[1]

        _remove_cell_borders(label_cell)
        _set_cell_shading(label_cell, light_bg)
        _set_cell_margins(label_cell, top=40, bottom=40, left=80, right=80)
        _set_cell_width(label_cell, 2520)
        lp = label_cell.paragraphs[0]
        _set_spacing(lp, before=0, after=0)
        _bold_run(lp, skill.category, s_label.font_name, s_label.size_hp, _hex_to_rgb(s_label.color_hex))

        _remove_cell_borders(value_cell)
        _set_cell_margins(value_cell, top=40, bottom=40, left=80, right=80)
        _set_cell_width(value_cell, 7560)
        vp = value_cell.paragraphs[0]
        _set_spacing(vp, before=0, after=0)
        _plain_run(vp, skill.items, s_value.font_name, s_value.size_hp, _hex_to_rgb(s_value.color_hex))

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(ctx.page.text_width))
    tblW.set(qn("w:type"), "dxa")
    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblPr.append(tblW)

    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tblBorders.append(el)
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    tblPr.append(tblBorders)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def render_docx(
    ir: ResumeIR,
    output_path: str,
    config: TemplateConfig | None = None,
) -> None:
    """Render *ir* to a styled DOCX file at *output_path*."""
    config = config or TemplateConfig()
    page = resolve_docx_page(config)

    ctx = SimpleNamespace(config=config, page=page)

    doc = Document()

    # -- Default font for the document --
    s_body = resolve_docx_style("body", config)
    style = doc.styles["Normal"]
    font = style.font
    font.name = s_body.font_name
    font.size = Pt(s_body.size_hp * HP)

    # -- Page setup --
    section = doc.sections[0]
    section.page_width = Twips(page.width)
    section.page_height = Twips(page.height)
    section.top_margin = Twips(page.top)
    section.bottom_margin = Twips(page.bottom)
    section.left_margin = Twips(page.left)
    section.right_margin = Twips(page.right)

    # -- Register bullet numbering --
    num_id = _configure_bullet_numbering(doc)

    # ===================================================================
    # HEADER — Name, Title, Contact
    # ===================================================================
    s_name = resolve_docx_style("name", config)
    s_title = resolve_docx_style("title", config)
    s_contact = resolve_docx_style("contact", config)

    if doc.paragraphs:
        p_name = doc.paragraphs[0]
    else:
        p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p_name, after=0)
    _bold_run(p_name, ir.header.name, s_name.font_name, s_name.size_hp, _hex_to_rgb(s_name.color_hex))

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p_title, before=40, after=60)
    _plain_run(p_title, ir.header.title, s_title.font_name, s_title.size_hp, _hex_to_rgb(s_title.color_hex))

    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(p_contact, after=100)

    muted_rgb = _hex_to_rgb(config.colors.muted)
    accent_rgb = _hex_to_rgb(config.colors.accent)
    contact_parts = [
        (ir.header.location, muted_rgb),
        ("  |  ", muted_rgb),
        (ir.header.email, muted_rgb),
        ("  |  ", muted_rgb),
        (ir.header.linkedin, accent_rgb),
        ("  |  ", muted_rgb),
        (ir.header.github, accent_rgb),
    ]
    for text, color in contact_parts:
        _plain_run(p_contact, text, s_contact.font_name, s_contact.size_hp, color)

    # ===================================================================
    # PROFESSIONAL SUMMARY
    # ===================================================================
    _add_section_heading(doc, "Professional Summary", ctx)

    p_summary = doc.add_paragraph()
    _set_spacing(p_summary, before=60, after=60)
    _add_parsed_runs(p_summary, ir.summary.paragraph, s_body.font_name, s_body.size_hp)

    for sb in ir.summary.bullets:
        _add_bullet(doc, f"**{sb.label}:** {sb.text}", num_id, ctx)

    # ===================================================================
    # SKILLS
    # ===================================================================
    _add_section_heading(doc, "Skills", ctx)
    _add_skills_table(doc, ir.skills, ctx)

    # ===================================================================
    # PROFESSIONAL EXPERIENCE
    # ===================================================================
    _add_section_heading(doc, "Professional Experience", ctx)

    s_company_desc = resolve_docx_style("company_desc", config)

    for company in ir.experience:
        _add_subheading(doc, company.company, company.location, ctx)

        if company.description:
            p_desc = doc.add_paragraph()
            _set_spacing(p_desc, before=0, after=20)
            _add_tab_stop_right(p_desc, page.text_width)
            _italic_run(p_desc, company.description, s_company_desc.font_name, s_company_desc.size_hp, _hex_to_rgb(s_company_desc.color_hex))
            r = _plain_run(p_desc, f"\t{company.dates}", s_company_desc.font_name, s_company_desc.size_hp, muted_rgb)
            r.italic = True

        for role in company.roles:
            _add_role_title(doc, role.title, role.dates, ctx)
            if role.description:
                _add_role_description(doc, role.description, ctx)
            for bullet in role.bullets:
                _add_bullet(doc, _format_bullet_text(bullet), num_id, ctx)

    # ===================================================================
    # PERSONAL PROJECTS
    # ===================================================================
    _add_section_heading(doc, "Personal Projects", ctx)

    s_proj = resolve_docx_style("project_title", config)
    s_pdesc = resolve_docx_style("project_desc", config)

    for project in ir.projects:
        p_proj = doc.add_paragraph()
        _set_spacing(p_proj, before=100, after=40)
        _bold_run(p_proj, project.name, s_proj.font_name, s_proj.size_hp)
        _plain_run(p_proj, " — ", s_proj.font_name, s_proj.size_hp)
        display_url = project.url.replace("https://", "").replace("http://", "")
        _plain_run(p_proj, display_url, s_proj.font_name, s_proj.size_hp, accent_rgb)

        p_pdesc = doc.add_paragraph()
        _set_spacing(p_pdesc, before=0, after=80)
        _add_parsed_runs(p_pdesc, project.description, s_pdesc.font_name, s_pdesc.size_hp)

    # ===================================================================
    # EDUCATION
    # ===================================================================
    _add_section_heading(doc, "Education", ctx)

    s_edu = resolve_docx_style("education", config)
    p_edu = doc.add_paragraph()
    _set_spacing(p_edu, before=60, after=0)
    _bold_run(p_edu, ir.education.degree, s_edu.font_name, s_edu.size_hp)
    _plain_run(p_edu, f" | {ir.education.institution}", s_edu.font_name, s_edu.size_hp)

    # ===================================================================
    # Save
    # ===================================================================
    doc.save(output_path)
